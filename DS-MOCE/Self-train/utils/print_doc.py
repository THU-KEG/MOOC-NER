# -*- coding:utf-8 -*-
import logging
import numpy as np
import torch
import docx
from torch.utils.data import DataLoader, SequentialSampler
from torch.utils.data.distributed import DistributedSampler
from tqdm import tqdm
from seqeval.metrics import precision_score
from seqeval.metrics import recall_score
from seqeval.metrics import f1_score

from utils.data_utils import load_and_cache_examples, tag_to_id
from flashtool import Logger
# logger = logging.getLogger(__name__)
# logging.basicConfig(
#     format="%(asctime)s - %(levelname)s - %(name)s -   %(message)s",
#     datefmt="%m/%d/%Y %H:%M:%S",
#     level=logging.INFO if args.local_rank in [-1, 0] else logging.WARN,
# )
# logging_fh = logging.FileHandler(os.path.join(args.output_dir, 'log.txt'))
# logging_fh.setLevel(logging.DEBUG)
# logger.addHandler(logging_fh)
# logger.warning(
#     "Process rank: %s, device: %s, n_gpu: %s, distributed training: %s, 16-bits training: %s",
#     args.local_rank,
#     device,
#     args.n_gpu,
#     bool(args.local_rank != -1),
#     args.fp16,
# )
def evaluate(args, model, tokenizer, labels, pad_token_label_id, best, mode, logger, prefix="", verbose=True):
    
    eval_dataset = load_and_cache_examples(args, tokenizer, labels, pad_token_label_id, mode=mode)

    args.eval_batch_size = args.per_gpu_eval_batch_size * max(1, args.n_gpu)
    eval_sampler = SequentialSampler(eval_dataset) if args.local_rank == -1 else DistributedSampler(eval_dataset)
    eval_dataloader = DataLoader(eval_dataset, sampler=eval_sampler, batch_size=args.eval_batch_size)

    # multi-gpu evaluate
    # if args.n_gpu > 1:
    #     model = torch.nn.DataParallel(model)

    logger.info("***** Running evaluation %s *****", prefix)
    if verbose:
        logger.info("  Num examples = %d", len(eval_dataset))
        logger.info("  Batch size = %d", args.eval_batch_size)
    eval_loss = 0.0
    nb_eval_steps = 0
    preds = None
    out_label_ids = None
    input_ids = None

    model.eval()
    for batch in tqdm(eval_dataloader, desc="Evaluating"):
        batch = tuple(t.to(args.device) for t in batch)

        with torch.no_grad():
            inputs = {"input_ids": batch[0], "attention_mask": batch[1], "labels": {"pseudo": batch[3]}}
            inputs['token_type_ids']=batch[2]
            inputs['field_embeds']=batch[4]
            outputs = model(**inputs)
            tmp_eval_loss_dict, logits = outputs[:2]
            tmp_eval_loss = tmp_eval_loss_dict["pseudo"]

            if args.n_gpu > 1:
                tmp_eval_loss = tmp_eval_loss.mean()

            eval_loss += tmp_eval_loss.item()
        nb_eval_steps += 1
        if preds is None:
            preds = logits.detach().cpu().numpy()
            out_label_ids = inputs["labels"]["pseudo"].detach().cpu().numpy()
            input_ids = inputs["input_ids"].detach().cpu().numpy()
        else:
            preds = np.append(preds, logits.detach().cpu().numpy(), axis=0)
            out_label_ids = np.append(out_label_ids, inputs["labels"]["pseudo"].detach().cpu().numpy(), axis=0)
            input_ids = np.append(input_ids,inputs['input_ids'].detach().cpu().numpy(),axis=0)    
    
    eval_loss = eval_loss / nb_eval_steps
    # print(preds)
    preds = np.argmax(preds, axis=2)

    label_map = {i: label for i, label in enumerate(labels)}
    preds_list = [[] for _ in range(out_label_ids.shape[0])]
    out_list =[[] for _ in range(out_label_ids.shape[0])] 
    #out_id_list = [[] for _ in range(out_label_ids.shape[0])]
    #preds_id_list = [[] for _ in range(out_label_ids.shape[0])]

    out_id_list = [[] for _ in range(out_label_ids.shape[0])]
    preds_id_list = [[] for _ in range(out_label_ids.shape[0])]
    input_id_list = [[] for _ in range(out_label_ids.shape[0])]

    for i in range(out_label_ids.shape[0]):
        for j in range(out_label_ids.shape[1]):
            if out_label_ids[i, j] != pad_token_label_id:
                preds_list[i].append(label_map[preds[i][j]])
                out_list[i].append(label_map[out_label_ids[i][j]])
                # preds_id_list[i].append(preds[i][j])
                out_id_list[i].append(out_label_ids[i][j])
                preds_id_list[i].append(preds[i][j])
                input_id_list[i].append(input_ids[i][j])


    p   = precision_score(out_list,preds_list)
    r   = recall_score(out_list,preds_list)
    new_F = f1_score(out_list,preds_list)

    is_updated = False
    if new_F > best[-1]:
        best = [p, r, new_F]
        is_updated = True

    results = {
       "loss": eval_loss,
       "precision": p,
       "recall": r,
       "f1": new_F,
       "best_precision": best[0],
       "best_recall":best[1],
       "best_f1": best[-1]
    }

    logger.info("***** Eval results %s *****", prefix)
    for key in sorted(results.keys()):
        logger.info("  %s = %s", key, str(results[key]))

    # write to doc:
    doc_name = 'SCDL_embeds'
    document = docx.Document()
    
    for per_content_id,per_ground_truth_ids,per_preds_ids in zip(input_id_list,out_id_list,preds_id_list):

        this_paragraph = document.add_paragraph()
        for index_id,index_truth,index_preds in zip(per_content_id,per_ground_truth_ids,per_preds_ids):
            # decode:
            this_word = tokenizer.decode(int(index_id))
            word = this_paragraph.add_run(this_word)
            if index_truth!=0 and index_preds!=0:
                word.font.highlight_color = docx.enum.text.WD_COLOR.RED
            elif index_truth!=0:
                word.font.highlight_color = docx.enum.text.WD_COLOR.YELLOW
            elif index_preds!=0:
                word.font.highlight_color = docx.enum.text.WD_COLOR.GRAY_25
        document.add_paragraph()
        
    
    document.save('{}.docx'.format(doc_name))
    

    return results, preds_list, best, is_updated
